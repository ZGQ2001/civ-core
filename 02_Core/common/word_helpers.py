"""Word 文档（python-docx）通用操作。

只放纯函数 / 纯逻辑，不启动 Word 应用、不弹窗。
"""

import re
from collections.abc import Callable

from docx import Document

from common.patterns import FIG_PATTERN
from common.types import PhotoPair


def open_first_table(doc_path: str):
    """打开 Word 文档，返回 (Document, first_table)。没有表格时抛错。"""
    doc = Document(doc_path)
    if not doc.tables:
        raise ValueError(f"❌ Word 文档没有任何表格: {doc_path}")
    return doc, doc.tables[0]


def scan_photo_pairs(
    doc_path: str, valid_nums: set | None = None
) -> tuple[dict[int, PhotoPair], list[PhotoPair]]:
    """扫描"上图下注"风格的表格，返回 ({num: PhotoPair}, [未匹配的 PhotoPair])。

    表格结构假设：偶数行（0、2、4…）放图，奇数行（1、3、5…）放对应题注。
    valid_nums 用来区分"已知排序的"和"未排序的"（None = 全都视为已匹配）。
    """
    print(f"🔍 扫描 Word 表格: {doc_path}")
    _, table = open_first_table(doc_path)
    total_rows = len(table.rows)
    print(f"📑 表格共 {total_rows} 行，开始解析...")

    matched: dict[int, PhotoPair] = {}
    unmatched: list[PhotoPair] = []

    for i in range(0, total_rows, 2):
        if i + 1 >= total_rows:
            break
        if i % 20 == 0:
            print(f"   ↳ 解析进度: {i}/{total_rows}")

        img_row = table.rows[i]
        txt_row = table.rows[i + 1]

        for j in range(len(img_row.cells)):
            text = txt_row.cells[j].text.strip()
            if not text:
                continue
            m = FIG_PATTERN.search(text)
            if not m:
                continue
            num = int(m.group(1))
            pair = PhotoPair(
                num=num, img_row_idx=i, txt_row_idx=i + 1, img_col_idx=j, txt_col_idx=j
            )
            if valid_nums is None or num in valid_nums:
                matched[num] = pair
            else:
                unmatched.append(pair)

    print(f"✅ 解析完成：匹配 {len(matched)} 个，未匹配 {len(unmatched)} 个")
    return matched, unmatched


def build_caption_renumber_mapping(doc_path: str) -> dict[int, int]:
    """按"行优先"扫描已排序文档的题注行，生成 {旧编号: 新编号}（新编号从 1 起）。

    专门给 renumber 工具用 —— sort_photos 输出的表格题注行是 1、3、5…（0-indexed）。
    """
    _, table = open_first_table(doc_path)
    mapping: dict[int, int] = {}
    new_num = 1

    for row_idx in range(1, len(table.rows), 2):
        row = table.rows[row_idx]
        seen_tc = set()  # 合并单元格去重
        for cell in row.cells:
            if cell._tc in seen_tc:
                continue
            seen_tc.add(cell._tc)

            text = cell.text.strip()
            if not text:
                continue
            m = FIG_PATTERN.search(text)
            if not m:
                continue
            old = int(m.group(1))
            if old in mapping:
                print(f"   ⚠️ 编号 {old} 重复出现，已忽略后续映射")
                continue
            mapping[old] = new_num
            new_num += 1

    print(f"✅ 已构建 {len(mapping)} 条编号映射 (1 → {new_num - 1})")
    preview = list(mapping.items())[:8]
    print("   预览（前 8 条 旧→新）:", "  ".join(f"{o}→{n}" for o, n in preview))
    return mapping


def make_caption_substitutor(mapping: dict[int, int]) -> tuple[Callable[[str], str], list[int]]:
    """生成一个"图 N → 图 mapping[N]"的字符串替换函数；保留原始空格前缀。

    返回 (apply, unmatched_log)：unmatched_log 在每次 apply 调用后会追加遇到但没有映射的旧编号。
    """
    unmatched: list[int] = []

    def sub(m: re.Match) -> str:
        old = int(m.group(1))
        new = mapping.get(old)
        if new is None:
            unmatched.append(old)
            return m.group(0)
        # 保留 "图" + 中间空白，只换数字部分
        prefix = m.group(0)[: m.start(1) - m.start(0)]
        return f"{prefix}{new}"

    def apply(text: str) -> str:
        return FIG_PATTERN.sub(sub, text)

    return apply, unmatched


def replace_in_caption_rows(
    doc_path: str, mapping: dict[int, int], output_path: str
) -> tuple[int, int, list[int]]:
    """改写第一个表格的题注行，保存到 output_path。

    返回 (run_级替换数, 段落级回退数, 找不到映射的旧编号列表)。
    保留每段格式：优先逐 run 替换；只有"图 N"被拆到多个 run 时才整段重写第一个 run。
    """
    doc, table = open_first_table(doc_path)
    apply, unmatched = make_caption_substitutor(mapping)

    run_count = 0
    fallback_count = 0

    for row_idx in range(1, len(table.rows), 2):
        for cell in table.rows[row_idx].cells:
            for paragraph in cell.paragraphs:
                touched = False
                for run in paragraph.runs:
                    if FIG_PATTERN.search(run.text):
                        new_text = apply(run.text)
                        if new_text != run.text:
                            run.text = new_text
                            run_count += 1
                            touched = True
                # run 内没匹配但段落整体有 → "图 N"被拆到多个 run，整段回退
                if not touched and FIG_PATTERN.search(paragraph.text):
                    full_new = apply(paragraph.text)
                    if full_new != paragraph.text and paragraph.runs:
                        paragraph.runs[0].text = full_new
                        for r in paragraph.runs[1:]:
                            r.text = ""
                        fallback_count += 1

    doc.save(output_path)
    return run_count, fallback_count, unmatched
