"""提取 docx 文本到 txt"""
from pathlib import Path
from docx import Document

src_dir = Path(r"D:\CodeProjects\civ-core\data\training_materials")

for fpath in sorted(src_dir.glob("*.docx")):
    name = fpath.name
    print(f"Processing: {name}")
    doc = Document(str(fpath))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    # Also extract tables
    for table in doc.tables:
        text += "\n---TABLE---\n"
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            text += " | ".join(cells) + "\n"
    out_path = src_dir / f"{fpath.stem}.txt"
    out_path.write_text(text, encoding="utf-8")
    print(f"  -> {out_path} ({len(text)} chars)")

print("\nDone.")
