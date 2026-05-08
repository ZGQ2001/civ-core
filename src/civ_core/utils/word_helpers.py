"""Word 文档（python-docx）通用操作（无 UI 依赖）。

工程规范落地：
  ✓ 全开类型注解
  ✓ 不再返回裸 tuple，返回 PhotoScanResult / CaptionRenumberMapping / CaptionReplaceResult
  ✓ print → logger
  ✓ raise IOReadError 取代 raise ValueError，UI 层可捕获后用 InfoBar 提示
"""

from __future__ import annotations

import re
from collections.abc import Callable
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table

from civ_core.models.schema import (
    CaptionRenumberMapping,
    CaptionReplaceResult,
    IOReadError,
    PhotoPair,
    PhotoScanResult,
)
from civ_core.utils.logger import get_logger
from civ_core.utils.patterns import FIG_PATTERN

log = get_logger(__name__)


# ──────────────────────────────────────────────────────────────────
# 1. 文档 / 表格打开
# ──────────────────────────────────────────────────────────────────
def open_first_table(doc_path: Path | str) -> tuple[DocxDocument, Table]:
    """打开 Word 文档，返回 (Document, first_table)。没有表格时抛 IOReadError。"""
    path = Path(doc_path)
    if not path.is_file():
        raise IOReadError(f"Word 文档不存在: {path}", hint=f"请确认路径正确：{path}")

    try:
        doc = Document(str(path))
    except Exception as e:
        raise IOReadError(f"无法打开 Word 文档 {path.name}: {e}") from e

    if not doc.tables:
        raise IOReadError(
            f"Word 文档没有任何表格: {path.name}",
            hint="请确认文档里至少包含一个表格再来运行该工具。",
        )
    return doc, doc.tables[0]


# ──────────────────────────────────────────────────────────────────
# 2. 扫描「上图下注」表格 → PhotoScanResult
# ──────────────────────────────────────────────────────────────────
def scan_photo_pairs(
    doc_path: Path | str,
    valid_nums: set[int] | None = None,
) -> PhotoScanResult:
    """扫描「上图下注」风格的表格。

    表格结构假设：偶数行 (0, 2, 4...) 放图，奇数行 (1, 3, 5...) 放对应题注。
    valid_nums 用来区分「已知排序的」和「未排序的」（None = 全都视为已匹配）。
    """
    path = Path(doc_path)
    log.info("扫描 Word 表格: %s", path.name)
    _, table = open_first_table(path)
    total_rows = len(table.rows)
    log.info("表格共 %d 行，开始解析...", total_rows)

    matched: dict[int, PhotoPair] = {}
    unmatched: list[PhotoPair] = []

    for i in range(0, total_rows, 2):
        if i + 1 >= total_rows:
            break
        if i and i % 20 == 0:
            log.debug("解析进度 %d/%d", i, total_rows)

        img_row = table.rows[i]
        txt_row = table.rows[i + 1]

        for j in range(len(img_row.cells)):
            text = txt_row.cells[j].text.strip()
            if not text:
                continue
            m = FIG_PATTERN.search(text)
            if not m:
                continue
            num = int(m.group(1))
            pair = PhotoPair(
                num=num,
                img_row_idx=i,
                txt_row_idx=i + 1,
                img_col_idx=j,
                txt_col_idx=j,
            )
            if valid_nums is None or num in valid_nums:
                matched[num] = pair
            else:
                unmatched.append(pair)

    log.info("解析完成：匹配 %d 个，未匹配 %d 个", len(matched), len(unmatched))
    return PhotoScanResult(matched=matched, unmatched=unmatched, total_rows=total_rows)


# ──────────────────────────────────────────────────────────────────
# 3. 题注重编号映射 → CaptionRenumberMapping
# ──────────────────────────────────────────────────────────────────
def build_caption_renumber_mapping(doc_path: Path | str) -> CaptionRenumberMapping:
    """按「行优先」扫描已排序文档的题注行，生成 {旧编号: 新编号}（新编号从 1 起）。

    专门给 renumber 工具用 —— sort_photos 输出的表格题注行是 1、3、5...（0-indexed）。
    """
    _, table = open_first_table(doc_path)
    mapping: dict[int, int] = {}
    duplicates: list[int] = []
    new_num = 1

    for row_idx in range(1, len(table.rows), 2):
        row = table.rows[row_idx]
        seen_tc = set()  # 合并单元格去重
        for cell in row.cells:
            if cell._tc in seen_tc:
                continue
            seen_tc.add(cell._tc)

            text = cell.text.strip()
            if not text:
                continue
            m = FIG_PATTERN.search(text)
            if not m:
                continue
            old = int(m.group(1))
            if old in mapping:
                log.warning("编号 %d 重复出现，已忽略后续映射", old)
                duplicates.append(old)
                continue
            mapping[old] = new_num
            new_num += 1

    log.info("已构建 %d 条编号映射 (1 → %d)", len(mapping), new_num - 1)
    if mapping:
        preview = list(mapping.items())[:8]
        log.debug("映射预览（前 8 条 旧→新）: %s", "  ".join(f"{o}→{n}" for o, n in preview))
    return CaptionRenumberMapping(mapping=mapping, duplicates=duplicates)


# ──────────────────────────────────────────────────────────────────
# 4. 题注字符串替换器（生成 closure）
# ──────────────────────────────────────────────────────────────────
def make_caption_substitutor(
    mapping: dict[int, int],
) -> tuple[Callable[[str], str], list[int]]:
    """生成「图 N → 图 mapping[N]」的字符串替换函数；保留原始空格前缀。

    返回 (apply, unmatched_log)：unmatched_log 在每次 apply 调用后会追加
    遇到但没有映射的旧编号 —— 同一份 list，调用方观察即可。
    """
    unmatched: list[int] = []

    def _sub(m: re.Match) -> str:
        old = int(m.group(1))
        new = mapping.get(old)
        if new is None:
            unmatched.append(old)
            return m.group(0)
        prefix = m.group(0)[: m.start(1) - m.start(0)]
        return f"{prefix}{new}"

    def apply(text: str) -> str:
        return FIG_PATTERN.sub(_sub, text)

    return apply, unmatched


# ──────────────────────────────────────────────────────────────────
# 5. 改写题注行 → CaptionReplaceResult
# ──────────────────────────────────────────────────────────────────
def replace_in_caption_rows(
    doc_path: Path | str,
    mapping: dict[int, int],
    output_path: Path | str,
) -> CaptionReplaceResult:
    """改写第一个表格的题注行，保存到 output_path。

    保留每段格式：优先逐 run 替换；只有「图 N」被拆到多个 run 时才整段重写第一个 run。
    """
    doc, table = open_first_table(doc_path)
    apply, unmatched = make_caption_substitutor(mapping)

    run_count = 0
    fallback_count = 0

    for row_idx in range(1, len(table.rows), 2):
        for cell in table.rows[row_idx].cells:
            for paragraph in cell.paragraphs:
                touched = False
                for run in paragraph.runs:
                    if FIG_PATTERN.search(run.text):
                        new_text = apply(run.text)
                        if new_text != run.text:
                            run.text = new_text
                            run_count += 1
                            touched = True
                # run 内没匹配但段落整体有 → 「图 N」被拆到多个 run，整段回退
                if not touched and FIG_PATTERN.search(paragraph.text):
                    full_new = apply(paragraph.text)
                    if full_new != paragraph.text and paragraph.runs:
                        paragraph.runs[0].text = full_new
                        for r in paragraph.runs[1:]:
                            r.text = ""
                        fallback_count += 1

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    log.info(
        "题注替换完成: run-level=%d, paragraph-fallback=%d, unmatched=%d → %s",
        run_count,
        fallback_count,
        len(unmatched),
        out.name,
    )
    return CaptionReplaceResult(
        output_path=out,
        run_level_replacements=run_count,
        paragraph_fallbacks=fallback_count,
        unmatched_old_ids=unmatched,
    )
